"""
Export Service for NoteTaker
Supports PDF, Markdown, and DOCX export functionality
"""
import os
import io
from datetime import datetime
from typing import List, Dict, Optional
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
import zipfile
import tempfile

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not installed. DOCX export will be unavailable.")

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation"""
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            textColor=HexColor('#2C3E50'),
            alignment=TA_CENTER
        )
        
        # Note title style
        self.note_title_style = ParagraphStyle(
            'NoteTitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=HexColor('#34495E')
        )
        
        # Content style
        self.content_style = ParagraphStyle(
            'Content',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=20
        )
        
        # Metadata style
        self.meta_style = ParagraphStyle(
            'Metadata',
            parent=self.styles['Normal'],
            fontSize=9,
            textColor=HexColor('#7F8C8D'),
            spaceAfter=8
        )
    
    def export_to_markdown(self, notes: List[Dict], include_translations: bool = True) -> str:
        """Export notes to Markdown format"""
        markdown_content = []
        
        # Header
        markdown_content.append("# NoteTaker Export")
        markdown_content.append(f"*Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        markdown_content.append("")
        markdown_content.append(f"**Total Notes:** {len(notes)}")
        markdown_content.append("")
        markdown_content.append("---")
        markdown_content.append("")
        
        # Notes
        for i, note in enumerate(notes, 1):
            # Note header
            markdown_content.append(f"## {i}. {note.get('title', 'Untitled')}")
            markdown_content.append("")
            
            # Metadata
            created_at = note.get('created_at', '')
            updated_at = note.get('updated_at', '')
            if created_at:
                markdown_content.append(f"**Created:** {created_at}")
            if updated_at:
                markdown_content.append(f"**Last Updated:** {updated_at}")
            
            # Tags
            tags = note.get('auto_tags', [])
            if tags:
                tags_str = ', '.join([f"`{tag}`" for tag in tags])
                markdown_content.append(f"**Tags:** {tags_str}")
            
            markdown_content.append("")
            
            # Content
            content = note.get('content', '')
            if content:
                markdown_content.append("### Content")
                markdown_content.append(content)
                markdown_content.append("")
            
            # Translation (if available and requested)
            if include_translations:
                title_zh = note.get('title_zh')
                content_zh = note.get('content_zh')
                
                if title_zh or content_zh:
                    markdown_content.append("### Chinese Translation")
                    if title_zh:
                        markdown_content.append(f"**Title (中文):** {title_zh}")
                    if content_zh:
                        markdown_content.append(f"**Content (中文):** {content_zh}")
                    markdown_content.append("")
            
            markdown_content.append("---")
            markdown_content.append("")
        
        return "\n".join(markdown_content)
    
    def export_to_pdf(self, notes: List[Dict], include_translations: bool = True) -> bytes:
        """Export notes to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        # Title
        story.append(Paragraph("NoteTaker Export", self.title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        export_info = f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Notes: {len(notes)}"
        story.append(Paragraph(export_info, self.meta_style))
        story.append(Spacer(1, 20))
        
        # Notes
        for i, note in enumerate(notes, 1):
            # Note title
            title = note.get('title', 'Untitled')
            story.append(Paragraph(f"{i}. {title}", self.note_title_style))
            
            # Metadata
            created_at = note.get('created_at', '')
            updated_at = note.get('updated_at', '')
            
            meta_info = []
            if created_at:
                meta_info.append(f"Created: {created_at}")
            if updated_at:
                meta_info.append(f"Updated: {updated_at}")
            
            # Tags
            tags = note.get('auto_tags', [])
            if tags:
                meta_info.append(f"Tags: {', '.join(tags)}")
            
            if meta_info:
                story.append(Paragraph(" | ".join(meta_info), self.meta_style))
            
            story.append(Spacer(1, 8))
            
            # Content
            content = note.get('content', '')
            if content:
                # Clean content for PDF
                clean_content = content.replace('\n', '<br/>')
                story.append(Paragraph(clean_content, self.content_style))
            
            # Translation
            if include_translations:
                title_zh = note.get('title_zh')
                content_zh = note.get('content_zh')
                
                if title_zh or content_zh:
                    story.append(Spacer(1, 8))
                    story.append(Paragraph("<b>Chinese Translation:</b>", self.content_style))
                    
                    if title_zh:
                        story.append(Paragraph(f"<b>Title:</b> {title_zh}", self.content_style))
                    if content_zh:
                        clean_content_zh = content_zh.replace('\n', '<br/>')
                        story.append(Paragraph(f"<b>Content:</b> {clean_content_zh}", self.content_style))
            
            story.append(Spacer(1, 20))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_docx(self, notes: List[Dict], include_translations: bool = True) -> Optional[bytes]:
        """Export notes to DOCX format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('NoteTaker Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        export_info = doc.add_paragraph()
        export_info.add_run(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Notes: {len(notes)}")
        export_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Notes
        for i, note in enumerate(notes, 1):
            # Note title
            title = note.get('title', 'Untitled')
            note_heading = doc.add_heading(f"{i}. {title}", level=1)
            
            # Metadata paragraph
            meta_para = doc.add_paragraph()
            meta_para.style = 'Intense Quote'
            
            created_at = note.get('created_at', '')
            updated_at = note.get('updated_at', '')
            
            meta_info = []
            if created_at:
                meta_info.append(f"Created: {created_at}")
            if updated_at:
                meta_info.append(f"Updated: {updated_at}")
            
            # Tags
            tags = note.get('auto_tags', [])
            if tags:
                meta_info.append(f"Tags: {', '.join(tags)}")
            
            if meta_info:
                meta_para.add_run(" | ".join(meta_info))
            
            # Content
            content = note.get('content', '')
            if content:
                content_para = doc.add_paragraph(content)
                content_para.style = 'Normal'
            
            # Translation
            if include_translations:
                title_zh = note.get('title_zh')
                content_zh = note.get('content_zh')
                
                if title_zh or content_zh:
                    trans_heading = doc.add_heading('Chinese Translation', level=2)
                    
                    if title_zh:
                        title_para = doc.add_paragraph()
                        title_para.add_run('Title: ').bold = True
                        title_para.add_run(title_zh)
                    
                    if content_zh:
                        content_para = doc.add_paragraph()
                        content_para.add_run('Content: ').bold = True
                        content_para.add_run(content_zh)
            
            # Add space between notes
            doc.add_paragraph()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_multiple_formats(self, notes: List[Dict], formats: List[str], include_translations: bool = True) -> bytes:
        """Export notes in multiple formats and return as ZIP file"""
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            for format_type in formats:
                try:
                    if format_type.lower() == 'markdown':
                        content = self.export_to_markdown(notes, include_translations)
                        zip_file.writestr(f"notes_export_{timestamp}.md", content)
                    
                    elif format_type.lower() == 'pdf':
                        content = self.export_to_pdf(notes, include_translations)
                        zip_file.writestr(f"notes_export_{timestamp}.pdf", content)
                    
                    elif format_type.lower() == 'docx' and DOCX_AVAILABLE:
                        content = self.export_to_docx(notes, include_translations)
                        if content:
                            zip_file.writestr(f"notes_export_{timestamp}.docx", content)
                
                except Exception as e:
                    print(f"Error exporting {format_type}: {e}")
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()

# Initialize service instance
export_service = ExportService()